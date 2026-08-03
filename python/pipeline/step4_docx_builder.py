import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from config import DOCX_DIR, OUTPUT_DOCX
from utils import ensure_dirs
from document import load_document


def setup_standard_page_margins(section):
    """ Đặt lề trang A4 chuẩn (0.75 inch) để văn bản trôi chảy tự nhiên """
    section.page_width = Inches(8.27)    # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def set_table_borders_none(table):
    """ Xóa toàn bộ viền bảng để tạo Bảng Ẩn (Invisible Table) """
    tblPr = table._tbl.tblPr
    # Xóa viền cũ nếu có
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    # Thêm viền trong suốt
    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)


def get_alignment_enum(align_str):
    align_str = (align_str or "LEFT").upper()
    if align_str == "CENTER":
        return WD_ALIGN_PARAGRAPH.CENTER
    if align_str == "RIGHT":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if align_str == "JUSTIFY":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def add_formatted_paragraph(
    container,
    text,
    font_name="Times New Roman",
    font_size=11,
    bold=False,
    italic=False,
    alignment="LEFT",
    space_after=4,
    space_before=0
):
    """ Thêm đoạn văn bản Native với khoảng cách dòng mượt mà """
    p = container.add_paragraph()
    p.alignment = get_alignment_enum(alignment)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def main():
    ensure_dirs(DOCX_DIR)

    data = load_document()
    doc = Document()

    for page_index, page in enumerate(data.get("pages", [])):
        if page_index == 0:
            setup_standard_page_margins(doc.sections[0])
        else:
            doc.add_page_break()
            section = doc.sections[-1]
            setup_standard_page_margins(section)

        # 1. Chèn Logo / Hình ảnh trang (nếu có)
        for image in page.get("images", []):
            img_path = image.get("path") or image.get("image_path")
            if img_path and os.path.exists(img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(6)
                img_p.paragraph_format.space_after = Pt(6)
                try:
                    img_p.add_run().add_picture(img_path, width=Inches(2.5))
                except Exception:
                    pass

        # 2. Gom các dòng chữ thành Khối theo thứ tự đọc (Reading Order Flow)
        lines = page.get("lines", [])

        # Phân loại khối 2 cột song song vs khối 1 cột đơn
        lines_sorted = sorted(lines, key=lambda l: (l.get("order") or l.get("line_id") or 0))

        # Gom dòng song song trên cùng mức Y để tạo Bảng Ẩn (Borderless Table)
        grouped_rows = []
        current_row = []

        for line in lines_sorted:
            if not current_row:
                current_row.append(line)
            else:
                last_y = current_row[-1]["bbox"]["y_min"]
                curr_y = line["bbox"]["y_min"]
                # Nếu 2 dòng nằm trên cùng 1 dải ngang Y (song song 2 cột)
                if abs(curr_y - last_y) < 18:
                    current_row.append(line)
                else:
                    grouped_rows.append(current_row)
                    current_row = [line]

        if current_row:
            grouped_rows.append(current_row)

        # 3. Dựng văn bản trôi chảy (Flow Layout)
        for row in grouped_rows:
            # Trường hợp 1 dòng đơn (Paragraph chuẩn)
            if len(row) == 1:
                line = row[0]
                text = line.get("text", "").strip()
                if not text:
                    continue

                layout_type = line.get("layout_type", "text_box")
                is_title = (layout_type == "title")
                is_heading = (layout_type == "heading")

                font_size = 14 if is_title else (12 if is_heading else 11)
                bold = is_title or is_heading or line.get("style", {}).get("bold", False)

                add_formatted_paragraph(
                    container=doc,
                    text=text,
                    font_name="Times New Roman",
                    font_size=font_size,
                    bold=bold,
                    alignment=line.get("alignment", "LEFT"),
                    space_before=6 if (is_title or is_heading) else 0,
                    space_after=4
                )

            # Trường hợp 2+ dòng song song (Tự tạo Bảng Ẩn Không Viền)
            else:
                row_sorted = sorted(row, key=lambda l: l["bbox"]["x_min"])
                num_cols = len(row_sorted)

                table = doc.add_table(rows=1, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders_none(table)

                cell_width = Inches(6.7 / num_cols)

                for col_idx, line in enumerate(row_sorted):
                    cell = table.rows[0].cells[col_idx]
                    cell.width = cell_width
                    text = line.get("text", "").strip()
                    if text:
                        p = cell.paragraphs[0]
                        p.alignment = get_alignment_enum(line.get("alignment", "LEFT"))
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
                        
                        run = p.add_run(text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(OUTPUT_DOCX)
    print("✅ ĐÃ XUẤT THÀNH CÔNG FILE WORD NATIVE FLOW TẠI:", OUTPUT_DOCX)


if __name__ == "__main__":
    main()
