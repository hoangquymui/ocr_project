import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_JSON = os.path.join(BASE_DIR, "output_scan_ocr", "layout_processed_result.json")
OUTPUT_DOCX = os.path.join(BASE_DIR, "output_scan_ocr", "final_output.docx")

def group_lines_into_paragraphs(lines, width):
    """
    Thuật toán gộp các dòng rời rạc thành các đoạn văn (Paragraph) logic.
    Nếu các dòng liên tiếp đều căn LEFT và nằm gần nhau, chúng thuộc cùng 1 đoạn.
    Nếu dòng có thuộc tính CENTER hoặc RIGHT, nó đứng độc lập một mình.
    """
    paragraphs = []
    current_para = []

    for line in lines:
        # Nếu là dòng căn giữa hoặc căn phải -> Tạo thành đoạn văn độc lập ngay lập tức
        if line["alignment"] in ["CENTER", "RIGHT"]:
            if current_para:
                paragraphs.append({"alignment": "LEFT", "text": " ".join(current_para)})
                current_para = []
            paragraphs.append({"alignment": line["alignment"], "text": line["text"]})
        else:
            # Nếu là căn LEFT (Văn bản thường)
            # Kiểm tra xem dòng này có phải là bắt đầu của đoạn mới (thụt lề) hay không
            # Ở đây chúng ta tạm gộp các dòng liên tiếp, nếu khoảng cách dòng quá xa có thể tách sau
            current_para.append(line["text"])

    # Đóng đoạn văn cuối cùng nếu còn sót
    if current_para:
        paragraphs.append({"alignment": "LEFT", "text": " ".join(current_para)})

    return paragraphs

def create_docx_from_layout(layout_data):
    # Khởi tạo một file Word trống
    doc = Document()

    # Thiết lập lề trang tiêu chuẩn (Top, Bottom, Left, Right = 1 inch ~ 2.54 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Thiết lập Style mặc định cho toàn bộ file Word (Font Times New Roman, Size 13)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    for page in layout_data["pages"]:
        print(f"-> Đang ghi dữ liệu Trang {page['page']} vào file Word...")
        
        # Gộp dòng thành đoạn văn cho trang hiện tại
        paragraphs_data = group_lines_into_paragraphs(page["lines"], page["width"])
        
        for p_data in paragraphs_data:
            # Thêm đoạn văn vào Word
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)  # Khoảng cách dưới mỗi đoạn văn
            p.paragraph_format.line_spacing = 1.3   # Giãn dòng 1.3 lines
            
            # Áp dụng căn lề dựa trên kết quả của Step 2
            if p_data["alignment"] == "CENTER":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif p_data["alignment"] == "RIGHT":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            # Thêm text vào đoạn văn
            run = p.add_run(p_data["text"])
            
            # Một mẹo nhỏ cho văn bản pháp luật: Tiêu đề in hoa thường viết Bold
            if p_data["alignment"] == "CENTER" and p_data["text"].isupper():
                run.bold = True

        # Nếu có nhiều trang, thêm dấu ngắt trang (Page Break) giữa các trang
        if page['page'] < len(layout_data["pages"]):
            doc.add_page_break()

    # Lưu file Word
    doc.save(OUTPUT_DOCX)

def main():
    if not os.path.exists(INPUT_JSON):
        print(f"LỖI: Không tìm thấy file {INPUT_JSON}. Hãy chạy step2_layout_analysis.py trước!")
        return

    with open(INPUT_JSON, "r", encoding="utf-8") as f:
        layout_data = json.load(f)

    print("Bắt đầu quá trình dựng cấu trúc file Word...")
    create_docx_from_layout(layout_data)
    
    print("\n--- HOÀN THÀNH TOÀN BỘ PIPELINE OCR ---")
    print(f"File Word (.docx) đã được xuất thành công tại:\n{OUTPUT_DOCX}")

if __name__ == "__main__":
    main()