import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def setup_standard_page_margins(section):
    """ Đặt lề trang A4 chuẩn (0.75 inch) để văn bản trôi chảy tự nhiên """
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def set_table_borders_none(table):
    """ Xóa toàn bộ viền bảng (Invisible Table) """
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)


def set_table_borders_visible(table):
    """ Thiết lập đường kẻ viền bảng hiển thị nét mảnh rõ ràng (Table Grid) """
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)


def is_line_inside_table(line, tables):
    if not tables:
        return False
    lb = line["bbox"]
    l_cx = (lb["x_min"] + lb["x_max"]) / 2
    l_cy = (lb["y_min"] + lb["y_max"]) / 2
    for tbl in tables:
        tb = tbl["bbox"]
        if tb["x_min"] <= l_cx <= tb["x_max"] and tb["y_min"] <= l_cy <= tb["y_max"]:
            return True
    return False


def get_alignment_enum(align_str):
    align_str = (align_str or "LEFT").upper()
    if align_str == "CENTER":
        return WD_ALIGN_PARAGRAPH.CENTER
    if align_str == "RIGHT":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if align_str == "JUSTIFY":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def add_formatted_paragraph(
    container,
    text,
    font_name="Times New Roman",
    font_size=11,
    bold=False,
    italic=False,
    alignment="LEFT",
    space_after=4,
    space_before=0
):
    """ Thêm đoạn văn bản Native với khoảng cách dòng mượt mà """
    p = container.add_paragraph()
    p.alignment = get_alignment_enum(alignment)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15

    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def run_step4_docx_builder(document_data, output_path):
    """
    Bước 4: Dựng file Word Reflowable Flow Layout kết hợp Bảng Ẩn và Đường Kẻ Bảng (Table Grid)
    """
    doc = Document()

    for idx, page in enumerate(document_data.get("pages", [])):
        if idx == 0:
            setup_standard_page_margins(doc.sections[0])
        else:
            doc.add_page_break()
            section = doc.sections[-1]
            setup_standard_page_margins(section)

        # 1. Chèn Logo / Hình ảnh trang (nếu có)
        for image in page.get("images", []):
            img_path = image.get("path") or image.get("image_path")
            if img_path and os.path.exists(img_path):
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(6)
                img_p.paragraph_format.space_after = Pt(6)
                try:
                    img_p.add_run().add_picture(img_path, width=Inches(2.5))
                except Exception:
                    pass

        lines = page.get("lines", [])
        tables_in_page = page.get("tables", [])

        if not lines:
            continue

        lines_sorted = sorted(lines, key=lambda l: l.get("order") or l.get("line_id") or 0)

        grouped_rows = []
        current_row = []

        for line in lines_sorted:
            if not current_row:
                current_row.append(line)
            else:
                last_y = current_row[-1]["bbox"]["y_min"]
                curr_y = line["bbox"]["y_min"]
                if abs(curr_y - last_y) < 18:
                    current_row.append(line)
                else:
                    grouped_rows.append(current_row)
                    current_row = [line]

        if current_row:
            grouped_rows.append(current_row)

        for row in grouped_rows:
            is_table_row = any(is_line_inside_table(l, tables_in_page) for l in row)

            if len(row) == 1 and not is_table_row:
                line = row[0]
                text = line.get("text", "").strip()
                if not text:
                    continue

                layout_type = line.get("layout_type", "text_box")
                is_title = (layout_type == "title")
                is_heading = (layout_type == "heading")

                font_size = 14 if is_title else (12 if is_heading else 11)
                bold = is_title or is_heading or line.get("style", {}).get("bold", False)

                add_formatted_paragraph(
                    container=doc,
                    text=text,
                    font_name="Times New Roman",
                    font_size=font_size,
                    bold=bold,
                    alignment=line.get("alignment", "LEFT"),
                    space_before=6 if (is_title or is_heading) else 0,
                    space_after=4
                )

            else:
                row_sorted = sorted(row, key=lambda l: l["bbox"]["x_min"])
                num_cols = len(row_sorted)

                table = doc.add_table(rows=1, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                if is_table_row:
                    set_table_borders_visible(table)
                else:
                    set_table_borders_none(table)

                cell_width = Inches(6.7 / num_cols)

                for col_idx, line in enumerate(row_sorted):
                    cell = table.rows[0].cells[col_idx]
                    cell.width = cell_width
                    text = line.get("text", "").strip()
                    if text:
                        p = cell.paragraphs[0]
                        p.alignment = get_alignment_enum(line.get("alignment", "LEFT"))
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        
                        run = p.add_run(text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10.5 if is_table_row else 11)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(output_path)
    return output_path
