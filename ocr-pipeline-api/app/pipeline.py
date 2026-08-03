import os
import cv2
import fitz
import numpy as np
from PIL import Image
import torch
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from app.vietocr_onnx_engine import VietOCRONNXPredictor


def remove_table_lines(binary):
    """ Loại bỏ toàn bộ kẻ bảng ngang và kẻ bảng dọc """
    kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    lines_h = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_h)

    kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    lines_v = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_v)

    table_lines = cv2.add(lines_h, lines_v)
    return cv2.subtract(binary, table_lines)


def detect_table_and_cover_regions_api(img_bgr):
    """ Phát hiện khu vực Bảng Biểu và Khung Bìa Trang từ ảnh BGR """
    page_h, page_w = img_bgr.shape[:2]
    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
    _, binary = cv2.threshold(gray, 240, 255, cv2.THRESH_BINARY_INV)

    kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1))
    lines_h = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_h)

    kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 50))
    lines_v = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel_v)

    table_grid = cv2.add(lines_h, lines_v)
    kernel_close = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 25))
    table_mask = cv2.morphologyEx(table_grid, cv2.MORPH_CLOSE, kernel_close, iterations=2)

    contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    tables = []
    cover_frames = []

    for idx, cnt in enumerate(contours, start=1):
        x, y, w, h = cv2.boundingRect(cnt)
        if w > page_w * 0.75 and h > page_h * 0.75:
            cover_frames.append({
                "frame_id": len(cover_frames) + 1,
                "type": "cover_frame",
                "bbox": {"x_min": float(x), "y_min": float(y), "x_max": float(x + w), "y_max": float(y + h)}
            })
        elif w >= 160 and h >= 60:
            tables.append({
                "table_id": len(tables) + 1,
                "type": "table_region",
                "bbox": {"x_min": float(x), "y_min": float(y), "x_max": float(x + w), "y_max": float(y + h)}
            })
    return tables, cover_frames


def is_line_inside_table(line, tables):
    if not tables:
        return False
    lb = line["bbox"]
    l_cx = (lb["x_min"] + lb["x_max"]) / 2
    l_cy = (lb["y_min"] + lb["y_max"]) / 2
    for tbl in tables:
        tb = tbl["bbox"]
        if tb["x_min"] <= l_cx <= tb["x_max"] and tb["y_min"] <= l_cy <= tb["y_max"]:
            return True
    return False


def process_layout_alignment(lines, width):
    if not lines:
        return []
    
    CENTER_X = width / 2
    lines.sort(key=lambda l: l["bbox"]["y_min"])
    
    final_ordered_lines = []
    current_row = []
    
    for line in lines:
        if not current_row:
            current_row.append(line)
        else:
            prev_line = current_row[-1]
            avg_height = ((prev_line["bbox"]["y_max"] - prev_line["bbox"]["y_min"]) + 
                          (line["bbox"]["y_max"] - line["bbox"]["y_min"])) / 2
            y_distance = abs(line["bbox"]["y_min"] - prev_line["bbox"]["y_min"])
            
            if y_distance < (avg_height * 0.4):
                current_row.append(line)
            else:
                if len(current_row) > 1:
                    current_row.sort(key=lambda l: l["bbox"]["x_min"])
                final_ordered_lines.extend(current_row)
                current_row = [line]
                
    if current_row:
        if len(current_row) > 1:
            current_row.sort(key=lambda l: l["bbox"]["x_min"])
        final_ordered_lines.extend(current_row)

    for line in final_ordered_lines:
        bbox = line["bbox"]
        line_width = bbox["x_max"] - bbox["x_min"]
        line_center_x = (bbox["x_min"] + bbox["x_max"]) / 2
        
        alignment = "LEFT"
        if abs(line_center_x - CENTER_X) < (width * 0.06) and line_width < (width * 0.7):
            alignment = "CENTER"
        elif bbox["x_min"] > (CENTER_X * 1.1) and line_width < (width * 0.45):
            alignment = "RIGHT"
            
        line["alignment"] = alignment
        
    return final_ordered_lines


def set_table_borders_none(table):
    """ Xóa toàn bộ viền bảng (Invisible Table cho văn bản 2 bên) """
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>\n'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)


def set_table_borders_visible(table):
    """ Thiết lập đường kẻ viền bảng hiển thị nét mảnh rõ ràng (Table Grid) """
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    borders_xml = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
        '</w:tblBorders>'
    )
    tblPr.append(borders_xml)


def get_alignment_enum(align_str):
    align_str = (align_str or "LEFT").upper()
    if align_str == "CENTER":
        return WD_ALIGN_PARAGRAPH.CENTER
    if align_str == "RIGHT":
        return WD_ALIGN_PARAGRAPH.RIGHT
    if align_str == "JUSTIFY":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def create_docx_document(pages_data, output_path):
    doc = Document()

    for idx, page in enumerate(pages_data):
        section = doc.sections[idx] if idx < len(doc.sections) else doc.add_section()
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        lines = page.get("lines", [])
        tables_in_page = page.get("tables", [])

        if not lines:
            continue

        lines_sorted = sorted(lines, key=lambda l: l.get("bbox", {}).get("y_min", 0))

        grouped_rows = []
        current_row = []

        for line in lines_sorted:
            if not current_row:
                current_row.append(line)
            else:
                last_y = current_row[-1]["bbox"]["y_min"]
                curr_y = line["bbox"]["y_min"]
                if abs(curr_y - last_y) < 18:
                    current_row.append(line)
                else:
                    grouped_rows.append(current_row)
                    current_row = [line]

        if current_row:
            grouped_rows.append(current_row)

        for row in grouped_rows:
            is_table_row = any(is_line_inside_table(l, tables_in_page) for l in row)

            if len(row) == 1 and not is_table_row:
                line = row[0]
                text = line.get("text", "").strip()
                if not text:
                    continue

                is_upper_title = text.isupper() and len(text.split()) < 15
                font_size = 14 if is_upper_title else 11
                bold = is_upper_title

                p = doc.add_paragraph()
                p.alignment = get_alignment_enum(line.get("alignment", "LEFT"))
                p.paragraph_format.space_before = Pt(4 if is_upper_title else 0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15

                run = p.add_run(text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run.font.color.rgb = RGBColor(0, 0, 0)

            else:
                row_sorted = sorted(row, key=lambda l: l["bbox"]["x_min"])
                num_cols = len(row_sorted)

                table = doc.add_table(rows=1, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                if is_table_row:
                    set_table_borders_visible(table)
                else:
                    set_table_borders_none(table)

                cell_width = Inches(6.7 / num_cols)

                for col_idx, line in enumerate(row_sorted):
                    cell = table.rows[0].cells[col_idx]
                    cell.width = cell_width
                    text = line.get("text", "").strip()
                    if text:
                        p = cell.paragraphs[0]
                        p.alignment = get_alignment_enum(line.get("alignment", "LEFT"))
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        
                        run = p.add_run(text)
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10.5 if is_table_row else 11)
                        run.font.color.rgb = RGBColor(0, 0, 0)

        if idx < len(pages_data) - 1:
            doc.add_page_break()

    doc.save(output_path)